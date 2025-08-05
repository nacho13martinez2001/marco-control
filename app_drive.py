
import os
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
print("API KEY:", os.getenv("OPENAI_API_KEY"))
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
assistant_id = os.getenv("ASSISTANT_ID")

from flask import Flask, render_template, request, jsonify, send_from_directory
import os
from datetime import datetime
from docx import Document
import re
from werkzeug.utils import secure_filename
from informes_utils import cargar_informes, guardar_informes
#from conversion_pdf import convertir_a_pdf
#from drive_utils_editable_google_docx import download_template, upload_document, list_templates

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "generated"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

TEMPLATE_LOCAL = os.path.join(app.config["UPLOAD_FOLDER"], "plantilla_temp.docx")
DOCX_LOCAL = os.path.join(app.config["UPLOAD_FOLDER"], "descargado.docx")
PDF_LOCAL = os.path.join(app.config["UPLOAD_FOLDER"], "generado.pdf")
ultimo_archivo_chat = ""
ultimo_contexto_estabilidad = {}  # Guarda PI y ORN temporalmente

@app.route("/pruebas/estabilidad")
def vista_prueba_estabilidad():
    return render_template("carga_excel_estabilidad.html")

@app.route("/", methods=["GET", "POST"])
def index():
    global ultimo_archivo_chat
    templates = []
#    templates = list_templates()
    informes = cargar_informes()
    doc_url = informes[0]["url"] if informes else ""
    ultimo_pdf_url = ""

    if request.method == "POST":
        if "reemplazar_docx" in request.form:
            archivo = request.files.get("archivo_editado")
            if archivo and archivo.filename.endswith(".docx"):
                archivo.save(DOCX_LOCAL)
                ultimo_archivo_chat = DOCX_LOCAL
                doc_url = upload_document(DOCX_LOCAL)["edit_url"]
                if informes:
                    informes[0]["url"] = doc_url
                guardar_informes(informes)
            return render_template("formulario_drive.html", templates=templates, informes=informes, doc_url=doc_url, pdf_url=ultimo_pdf_url, texto_editor="")

        elif "plantilla_id" in request.form:
            plantilla_id = request.form["plantilla_id"]
            entidad = request.form["entidad"]
            ejercicio = request.form["ejercicio"]
            fecha = request.form["fecha"]
            interventor = request.form["interventor"]
            res = request.form.get("resolucion", "Pendiente de Resultado")

            download_template(plantilla_id, TEMPLATE_LOCAL)
            doc = Document(TEMPLATE_LOCAL)
            for p in doc.paragraphs:
                p.text = p.text.replace("[entidad]", entidad)
                p.text = p.text.replace("[ejercicio]", ejercicio)
                p.text = p.text.replace("[fecha]", fecha)
                p.text = p.text.replace("[interventor]", interventor)
                p.text = p.text.replace("[res]", res)

            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            nombre = f"Informe_{entidad}_{ejercicio}_{timestamp}.docx".replace(" ", "_")
            doc.save(DOCX_LOCAL)
            ultimo_archivo_chat = DOCX_LOCAL
            urls = upload_document(DOCX_LOCAL)
            doc_url = urls["edit_url"]
            informes = [{
                "nombre": nombre,
                "url": doc_url,
                "local": DOCX_LOCAL
            }]
            guardar_informes(informes)

            texto_editor = ""
            if os.path.exists(DOCX_LOCAL):
                try:
                    doc = Document(DOCX_LOCAL)
                    texto_editor = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception as e:
                    texto_editor = f"⚠️ No se pudo leer el archivo: {e}"

            return render_template("formulario_drive.html", templates=templates, informes=informes, doc_url=doc_url, pdf_url=ultimo_pdf_url, texto_editor=texto_editor)

    texto_editor = ""
    if ultimo_archivo_chat and os.path.exists(ultimo_archivo_chat):
        try:
            doc = Document(ultimo_archivo_chat)
            texto_editor = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            texto_editor = f"⚠️ No se pudo leer el archivo: {e}"

    return render_template("formulario_drive.html", templates=templates, informes=informes, doc_url=doc_url, pdf_url=ultimo_pdf_url, texto_editor=texto_editor)

import pandas as pd
from flask import request, jsonify
from werkzeug.utils import secure_filename
import os
import fitz
from docx import Document

@app.route("/chat_upload", methods=["POST"])
def chat_upload():
    archivo = request.files.get("archivo_chat")
    if not archivo:
        return jsonify({"respuesta": "❌ No se recibió ningún archivo."}), 400

    nombre = secure_filename(archivo.filename)
    ruta = os.path.join(app.config["UPLOAD_FOLDER"], nombre)
    archivo.save(ruta)

    try:
        if archivo.filename.endswith(".docx"):
            doc = Document(ruta)
            texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return jsonify({"texto": texto})

        elif archivo.filename.endswith(".pdf"):
            doc = fitz.open(ruta)
            texto = "\n\n".join(page.get_text() for page in doc)
            return jsonify({"texto": texto})

        elif archivo.filename.endswith(".xlsx") or archivo.filename.endswith(".csv"):
            if archivo.filename.endswith(".csv"):
                import io
                archivo.stream.seek(0)
                df = pd.read_csv(io.StringIO(archivo.stream.read().decode("latin-1")))
                nombre_hoja = archivo.filename.rsplit(".", 1)[0]
                excel = {nombre_hoja: df}
            else:
                excel = pd.read_excel(ruta, sheet_name=None)

            hojas = list(excel.keys())
            contenido = {}

            for hoja, df in excel.items():
                for col in df.select_dtypes(include=["float", "float64"]).columns:
                    df[col] = df[col].round(2)
                contenido[hoja] = df.to_csv(index=False, sep="\t")

            texto = contenido[hojas[0]]

            return jsonify({
                "hojas": hojas,
                "contenido": contenido,
                "texto": texto
            })

        # ⚠️ Si el archivo tiene extensión no contemplada:
        return jsonify({"respuesta": "⚠️ Tipo de archivo no soportado."}), 415

    except Exception as e:
        return jsonify({
            "respuesta": f"⚠️ No se pudo procesar el archivo: {str(e)}",
            "texto": ""
        }), 500

@app.route("/chat", methods=["POST"])
def chat():
    try:
        data = request.get_json()
        user_messages = data.get("messages", [])

        if not user_messages:
            return jsonify({"respuesta": "⚠️ No se recibió ningún mensaje válido."}), 400

        last_user_message = user_messages[-1]["content"]

        # 1. Crear thread
        thread = client.beta.threads.create()

        # 2. Añadir mensaje del usuario al thread
        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=last_user_message
        )

        # 3. Lanzar ejecución del assistant
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id
        )

        # 4. Esperar a que termine
        import time
        while True:
            run_status = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id
            )
            if run_status.status == "completed":
                break
            elif run_status.status == "failed":
                return jsonify({"respuesta": "❌ Fallo al procesar con el Assistant."}), 500
            time.sleep(1)

        # 5. Recuperar la respuesta final
        messages = client.beta.threads.messages.list(thread_id=thread.id)
        for m in messages.data:
            if m.role == "assistant":
                return jsonify({"respuesta": m.content[0].text.value})

        return jsonify({"respuesta": "❌ No se obtuvo respuesta del Assistant."}), 500

    except Exception as e:
        return jsonify({"respuesta": f"❌ Error: {str(e)}"}), 500@app.route("/generated/<path:filename>")
def generated_files(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)
        
@app.route("/biblioteca")
def biblioteca_conocimiento():
    return send_from_directory("static", "biblioteca_conocimiento.html")

@app.route("/leer_pdf")
def leer_pdf():
    from flask import request
    import fitz  # PyMuPDF

    archivo = request.args.get("archivo", "")
    ruta_local = archivo.replace("/static/", "static/")  # Asegura acceso local

    try:
        doc = fitz.open(ruta_local)
        texto = "\n\n".join([page.get_text() for page in doc])
        return jsonify({"texto": texto})
    except Exception as e:
        return jsonify({"texto": f"❌ Error al leer PDF: {e}"})

@app.route("/guardar_ficha/<nombre>", methods=["POST"])
def guardar_ficha(nombre):
    import os, json
    ruta = os.path.join("static", "conocimiento", nombre)
    datos = request.get_json()

    try:
        with open(ruta, "w", encoding="utf-8") as f:
            json.dump(datos, f, indent=2, ensure_ascii=False)
        return jsonify({"mensaje": "Ficha actualizada correctamente."})
    except Exception as e:
        return jsonify({"mensaje": f"❌ Error al guardar: {str(e)}"}), 500

@app.route("/chat_extendido_ia", methods=["POST"])
def chat_extendido_ia():
    datos = request.get_json()
    texto = datos.get("texto", "").strip()
    if len(texto) > 4000:
        texto = texto[:4000]
    ingresos = datos.get("ingresos", "desconocido")
    gastos = datos.get("gastos", "desconocido")

    if not texto:
        return jsonify({"respuesta": "⚠️ No se ha proporcionado texto para analizar."})

    prompt = f"""Eres un interventor de administración local. Elabora un análisis técnico completo del siguiente informe, identificando observaciones relevantes, riesgos, equilibrio presupuestario, cumplimiento normativo y aspectos a mejorar.

Redacta observaciones con lenguaje técnico, justificadas y redactadas como si formaran parte de un informe oficial de fiscalización.

Datos aportados:
- Ingresos previstos: {ingresos} €
- Gastos previstos: {gastos} €

Informe:
{text}"""

    try:
        respuesta = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "Eres un experto en intervención y fiscalización presupuestaria local."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1200,
            temperature=0.4
        )
        texto_respuesta = respuesta.choices[0].message.content.strip()
        return jsonify({"respuesta": texto_respuesta})
    except Exception as e:
        return jsonify({"respuesta": f"❌ Error al contactar con OpenAI: {e}"}), 500

from flask import request, jsonify
from openai import OpenAI
@app.route("/estructura_ia", methods=["POST"])
def estructura_ia():
    datos = request.get_json()
    texto = datos.get("texto", "").strip()

    if not texto:
        return jsonify({"estructura": [], "error": "❌ No se recibió texto válido."}), 400

    prompt = (
        "Extrae del siguiente informe los apartados identificados con numeración romana. "
        "Devuélvelos *exclusivamente* como un array JSON con la forma: "
        "[{\"numeral\": \"I.\", \"titulo\": \"Nombre del apartado\"}, ...] "
        "Sin texto introductorio ni explicaciones.\n\n"
        f"{texto}"
    )

    try:
        respuesta = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "Eres un experto en estructuración de informes técnicos."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=800,
            temperature=0.2
        )

        contenido = respuesta.choices[0].message.content.strip()

        import json
        try:
            estructura = json.loads(contenido)
            return jsonify({"estructura": estructura})
        except json.JSONDecodeError:
            return jsonify({
                "estructura": [],
                "error": "⚠️ La IA no devolvió un JSON válido.",
                "respuesta": contenido
            }), 500

    except Exception as e:
        return jsonify({"estructura": [], "error": f"❌ Error interno: {e}"}), 500

@app.route("/formulario_aprobacion")
def formulario_aprobacion():
    import pyodbc
    conn = pyodbc.connect(
        "DRIVER={ODBC Driver 17 for SQL Server};"
        "SERVER=localhost\\SQLEXPRESS;"
        "DATABASE=SC_ENTIDAD1_DAT;"
        "Trusted_Connection=yes;"
    )
    cursor = conn.cursor()
    cursor.execute("""
        SELECT TOP 1 xdocumento
        FROM imp.back_pc_inftit
        WHERE xdocumento IS NOT NULL AND LEN(xdocumento) > 0
    """)
    row = cursor.fetchone()

    datos = {
        "ingresos_ayto_a2": row[0] if row else "",  # o row.xdocumento si usas .description
        "gastos_ayto_b2": "12345",
        "ingresos_consolidados_c2": "67890",
        "gastos_consolidados_d2": "3456",
        "orns_a2": "444",
        "drns_b2": "555",
        "capacidad_n_c2": "666",
        "capacidad_pct_d2": "777",
        "enf_anterior": "888",
        "tasa_incremento": "1.2%",
        "limite_gasto": "9999"
    }

    return render_template("formulario_aprobacion.html", datos=datos)

    import pyodbc

    # Conexión a la base de datos SQL Server
    conn = pyodbc.connect(
        "DRIVER={ODBC Driver 17 for SQL Server};"
        "SERVER=localhost\\SQLEXPRESS;"
        "DATABASE=SC_ENTIDAD1_DAT;"
        "Trusted_Connection=yes;"
    )
    cursor = conn.cursor()

    # Leer valores desde la tabla imp.back_pc_inftit
    cursor.execute("""
        SELECT xnombre, xdocumento
        FROM imp.back_pc_inftit
        WHERE xdocumento IS NOT NULL AND LEN(xdocumento) > 0
    """)
    registros = cursor.fetchall()

    # Mapeamos manualmente qué xnombre rellena qué campo del formulario
    datos = {
        "ingresos_ayto_a2": "",
        "gastos_ayto_b2": "",
        "ingresos_consolidados_c2": "",
        "gastos_consolidados_d2": "",
        "orns_a2": "",
        "drns_b2": "",
        "capacidad_n_c2": "",
        "capacidad_pct_d2": "",
        "enf_anterior": "",
        "tasa_incremento": "",
        "limite_gasto": ""
    }

    for nombre, valor in registros:
        if not nombre:
            continue
        nombre = nombre.lower()
        if "ingresos" in nombre and "ayuntamiento" in nombre:
            datos["ingresos_ayto_a2"] = valor
        elif "gastos" in nombre and "ayuntamiento" in nombre:
            datos["gastos_ayto_b2"] = valor
        elif "ingresos consolidados" in nombre:
            datos["ingresos_consolidados_c2"] = valor
        elif "gastos consolidados" in nombre:
            datos["gastos_consolidados_d2"] = valor
        elif "orns" in nombre:
            datos["orns_a2"] = valor
        elif "drns" in nombre:
            datos["drns_b2"] = valor
        elif "capacidad" in nombre and "nº" in nombre:
            datos["capacidad_n_c2"] = valor
        elif "capacidad" in nombre and "%" in nombre:
            datos["capacidad_pct_d2"] = valor
        elif "enf n-1" in nombre:
            datos["enf_anterior"] = valor
        elif "tasa incremento" in nombre:
            datos["tasa_incremento"] = valor
        elif "límite gasto" in nombre:
            datos["limite_gasto"] = valor

    return render_template("formulario_aprobacion.html", datos=datos)

    import pyodbc
    conn = pyodbc.connect(
        "DRIVER={ODBC Driver 17 for SQL Server};"
        "SERVER=localhost\\SQLEXPRESS;"
        "DATABASE=SC_ENTIDAD1_DAT;"
        "Trusted_Connection=yes;"
    )
    cursor = conn.cursor()
    cursor.execute("""
        SELECT xnombre, xdocumento
        FROM imp.back_pc_inftit
        WHERE xdocumento IS NOT NULL AND LEN(xdocumento) > 0
    """)
    registros = cursor.fetchall()

    datos = {
        "ingresos_ayto_a2": "",
        "gastos_ayto_b2": "",
        "ingresos_consolidados_c2": "",
        "gastos_consolidados_d2": "",
        "orns_a2": "",
        "drns_b2": "",
        "capacidad_n_c2": "",
        "capacidad_pct_d2": "",
        "enf_anterior": "",
        "tasa_incremento": "",
        "limite_gasto": ""
    }

    for nombre, valor in registros:
        if "ingresos ayuntamiento (a2)" in nombre.lower():
            datos["ingresos_ayto_a2"] = valor
        elif "gastos ayuntamiento (b2)" in nombre.lower():
            datos["gastos_ayto_b2"] = valor
        elif "ingresos consolidados (c2)" in nombre.lower():
            datos["ingresos_consolidados_c2"] = valor
        elif "gastos consolidados (d2)" in nombre.lower():
            datos["gastos_consolidados_d2"] = valor
        elif "orns" in nombre.lower():
            datos["orns_a2"] = valor
        elif "drns" in nombre.lower():
            datos["drns_b2"] = valor
        elif "capacidad financiación nº (c2)" in nombre.lower():
            datos["capacidad_n_c2"] = valor
        elif "capacidad financiación % (d2)" in nombre.lower():
            datos["capacidad_pct_d2"] = valor
        elif "enf n-1" in nombre.lower():
            datos["enf_anterior"] = valor
        elif "tasa incremento" in nombre.lower():
            datos["tasa_incremento"] = valor
        elif "límite gasto" in nombre.lower():
            datos["limite_gasto"] = valor

    return render_template("formulario_aprobacion.html", datos=datos)

if __name__ == "__main__":
    app.run(debug=True)